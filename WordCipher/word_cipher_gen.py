from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.shared import Cm
import numpy as np

# Unicodes are used the "encrypt" the text we want
unicodes_list = ["\u2582", "\u25a0", "\u25c6", "\u25b6", "\u259a",
                 "\u25ce", "\u2665", "\u2663", "\u25d3", "\u25f0",
                 "\u2602", "\u273f", "\u2606", "\u2691", "\u25e1",
                 "\u2702", "\u266b", "\u26cc"]

TEXT_SIZE = 16
SQUARE_SIZE = 28
LETTER_16_WIDTH = 0.2
MARGIN_WIDTH = 0.9

class WordCipherGen:
    def __init__(self):
        """
        Init the class
        """
        self.mapping_dict = None
        self.solutions_enc = None
        self.secret_enc = None
        self.doc = None

    def _new_doc(self):
        # Create a new doc:
        self.doc = Document()

    def encrypt_text(self, solutions, secret):
        """
        Creates an encryption for the solutions and secret text, replacing each unique char with one of our unicodes
        Sets the mapping key in self.mapping_dict, the solutions encoded in self.solutions_enc and the secret encoded in self.secret_enc
        :param solutions: the list of strings that are the solutions
        :param secret: The secret text we want to discover last
        :return:
        """
        unique_chars = set([l for l in "".join(solutions)])
        assert all(l in unique_chars for l in secret), Exception("Not all chars in secret are in the solutions")

        # Now create a random order of the unicodes:
        unicode_ord = np.random.choice(len(unicodes_list), len(unique_chars), replace=False)
        self.mapping_dict = {c:k for c,k in zip(unique_chars, unicode_ord)}

        # Now "encrypt" the solutions and the secret:
        self.solutions_enc = [[unicodes_list[self.mapping_dict[c]] for c in sol] for sol in solutions]
        self.secret_enc = [unicodes_list[self.mapping_dict[c]] for c in secret]

    def _set_font_size(self, paragraphs, font_size=20):
        for paragraph in paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(font_size)  # Set font size to 20

    def _get_width(self, txt, font_size):
        return MARGIN_WIDTH + (len(txt) * LETTER_16_WIDTH) * font_size / 16.

    # Function to create a table with 6 columns and add text in the first row
    def create_table_with_columns_and_text(self, question, ans_symbols):
        """
        Create a table in the doc to contain the question text, then the symbols for the solution beneath empty squares
        :param question: The question text
        :param ans_symbols: The answer symbols as a list of chars
        :return: The table created
        """
        n_cols = len(ans_symbols) + 1

        # Add a table, aligned to the right:
        table = self.doc.add_table(rows=2, cols=n_cols)
        table.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        table.allow_autofit = True  # Disable autofit

        # Set the width of the first five columns for a single character
        for i in range(n_cols-1):
            cell = table.cell(0, i)
            cell.text = "\u25A1"
            cell.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM  # Center align vertically
            cell.width = Cm(1)  # Adjust the width as needed
            self._set_font_size(cell.paragraphs, SQUARE_SIZE)

            cell = table.cell(1, i)
            cell.width = Cm(1)  # Adjust the width as needed
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP  # Center align vertically
            cell.text = ans_symbols[i]
            self._set_font_size(cell.paragraphs, TEXT_SIZE)

        # Set the width of the last column to fill the remaining space on the page
        cell = table.cell(0, n_cols-1)
        cell.width = Cm(self._get_width(question, TEXT_SIZE))  # Adjust the width as needed
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center align vertically
        cell.text = question
        self._set_font_size(cell.paragraphs, TEXT_SIZE)

        return table

    # Function to add an empty paragraph for spacing
    def add_empty_paragraph(self):
        p = self.doc.add_paragraph()
        run = p.add_run("")
        run.font.size = Pt(1)  # Set font size to the smallest value for an empty line

    def create_doc(self, sentences, secret):
        """
        Create the doc from the sentences and secret.
        :param sentences: A list of tuples of (question, answer). Answers are one word, and encrypted with the symbols
        :param secret: A tuple of (question, answer) like the sentences
        :return:
        """
        # Create the encrypted data for the sentences and secret answers:
        self.encrypt_text([s[1][::-1] for s in sentences], secret[1][::-1])

        # Create a new Word document
        self._new_doc()

        # Add questions with symbols to the document
        for i in range(len(sentences)):
            question = sentences[i][0]
            symbols = self.solutions_enc[i]
            self.create_table_with_columns_and_text(question, symbols)
            self.add_empty_paragraph()

        # Now add the last "secret" question:
        self.create_table_with_columns_and_text(secret[0], self.secret_enc)

    def save_doc(self, filename):
        # Save the document
        self.doc.save(filename)


if __name__ == '__main__':

    # Your input sentences and symbols
    sentences = [
        ("איזו עונה מתחילה בסוכות", "סתיו"),
        ("פרח גבוה שפורח עכשיו", "חצב"),
        ("בונים אותה בחג סוכות", "סוכה"),
        ("פרי עם המון גרעינים אדומים", "רימונ"),
        ("חג סוכות נקרא גם חג", "האסיפ")
    ]
    secret = (" יודעים שהתחיל החג לפי", "הכוכבימ")
    word_cipher = WordCipherGen()
    word_cipher.create_doc(sentences, secret)
    word_cipher.save_doc("/Users/talsegalov/Downloads/hebrew_questions.docx")

