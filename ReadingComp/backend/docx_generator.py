from docx import Document
from docx.shared import Inches
import tempfile

from backend.database.models import Story


class DocxGenerator:
    def generate(self, story: Story) -> str:
        doc = Document()

        # Add title
        doc.add_heading(story.title, level=1)

        # Add text and image with square wrap and right alignment
        if story.image_bytes:
            # Add text
            text_paragraph = doc.add_paragraph(story.content)
            
            # Add image with right alignment and square wrap
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_file:
                tmp_file.write(story.image_bytes)
                tmp_file_path = tmp_file.name
                run = text_paragraph.add_run()
                run.add_picture(tmp_file_path, width=Inches(3.0))
                
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.runs[-1].add_break()
        else:
            # Add text if no image is present
            doc.add_paragraph(story.content)

        # Add questions heading
        doc.add_heading("Questions", level=2)

        # Add questions
        yes_no_questions = [q for q in story.questions if q.type == 'yes_no']
        other_questions = [q for q in story.questions if q.type != 'yes_no']

        for idx, q in enumerate(other_questions, 1):
            if q.type == 'multiple_choice':
                doc.add_paragraph(f"{idx}. {q.question}")
                for i, opt in enumerate(q.options):
                    doc.add_paragraph(f"\t{chr(97 + i)}. {opt}")
            else:
                doc.add_paragraph(f"{idx}. {q.question}")
                doc.add_paragraph("Answer: ________________________")

        if yes_no_questions:
            table = doc.add_table(rows=len(yes_no_questions) + 1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Question"
            hdr_cells[1].text = "Yes"
            hdr_cells[2].text = "No"
            for i, q in enumerate(yes_no_questions, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = q.question
                row_cells[1].text = "□"
                row_cells[2].text = "□"

        # Save to temp file
        temp_path = tempfile.mktemp(suffix=".docx")
        doc.save(temp_path)
        return temp_path
